"""
Script d'indexation du corpus « création d'entreprise ».

Lit tous les PDF et DOCX des deux dossiers sources, découpe leur texte en
morceaux, les vectorise avec l'API OpenAI et les range dans ChromaDB.
Chaque morceau garde sa source (nom de fichier, page) et son type
(personnelle / societaire), ce qui permettra de filtrer la recherche.

À lancer depuis le dossier BackendTangaAI :
    python -m rag.ingest

À relancer à chaque fois que les documents changent (reconstruction complète).
"""
from pathlib import Path

import pdfplumber
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import chromadb
from chromadb.utils import embedding_functions

from rag.config import (
    SOURCES, CHROMA_DIR, COLLECTION_NAME, EMBEDDING_MODEL,
    CHUNK_SIZE, CHUNK_OVERLAP, OPENAI_API_KEY,
)


def lire_pdf(chemin: Path):
    """Extrait le texte d'un PDF, page par page.

    Retourne une liste de tuples (texte, numero_page). Les pages vides
    (souvent des pages scannées sans couche texte) sont ignorées.
    """
    pages = []
    with pdfplumber.open(chemin) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            texte = page.extract_text() or ""
            if texte.strip():
                pages.append((texte, i))
    return pages


def lire_docx(chemin: Path):
    """Extrait le texte d'un fichier Word (.docx) : paragraphes + tableaux.

    Retourne une liste avec un seul tuple (texte_complet, None) — un .docx
    n'a pas de pagination fiable, donc pas de numéro de page.
    """
    doc = Document(chemin)
    morceaux = [p.text for p in doc.paragraphs if p.text.strip()]
    # On récupère aussi le contenu des tableaux (fréquent dans les formulaires)
    for table in doc.tables:
        for row in table.rows:
            cellules = [c.text.strip() for c in row.cells if c.text.strip()]
            if cellules:
                morceaux.append(" | ".join(cellules))
    texte = "\n".join(morceaux)
    return [(texte, None)] if texte.strip() else []


def lire_document(chemin: Path):
    """Aiguille vers le bon lecteur selon l'extension du fichier."""
    ext = chemin.suffix.lower()
    if ext == ".pdf":
        return lire_pdf(chemin)
    if ext == ".docx":
        return lire_docx(chemin)
    return []  # autres extensions ignorées


def construire_index():
    """Construit (ou reconstruit) l'index vectoriel à partir des dossiers sources."""
    if not OPENAI_API_KEY:
        raise RuntimeError("OPENAI_API_KEY introuvable. Vérifiez votre fichier .env.")

    decoupeur = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    documents, metadatas, ids = [], [], []
    compteur_fichiers = 0

    # Parcours des deux dossiers (chaque dossier = un type d'entreprise)
    for type_entreprise, dossier in SOURCES.items():
        if not dossier.exists():
            print(f"[ATTENTION] Dossier introuvable : {dossier}")
            continue

        print(f"\nDossier « {type_entreprise} » : {dossier}")
        for chemin in sorted(dossier.iterdir()):
            if chemin.suffix.lower() not in (".pdf", ".docx"):
                continue

            unites = lire_document(chemin)
            if not unites:
                print(f"  [ATTENTION] Aucun texte extrait de {chemin.name} "
                      f"(PDF scanné ? il faudrait passer par l'OCR).")
                continue

            compteur_fichiers += 1
            n_morceaux = 0
            for texte_unite, page in unites:
                for j, morceau in enumerate(decoupeur.split_text(texte_unite)):
                    documents.append(morceau)
                    metadatas.append({
                        "type": type_entreprise,
                        "source": chemin.name,
                        "page": page if page is not None else -1,
                    })
                    page_label = page if page is not None else 0
                    ids.append(f"{type_entreprise}__{chemin.stem}__{page_label}__{j}")
                    n_morceaux += 1
            print(f"  + {chemin.name} → {n_morceaux} morceaux")

    if not documents:
        raise RuntimeError("Aucun document indexé. Vérifiez les chemins des dossiers sources.")

    print(f"\n{compteur_fichiers} fichiers lus, {len(documents)} morceaux à indexer.")

    # Connexion à ChromaDB (persistance locale) + embedding OpenAI
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    openai_ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=OPENAI_API_KEY,
        model_name=EMBEDDING_MODEL,
    )

    # Reconstruction propre : on supprime l'ancienne collection si elle existe
    try:
        client.delete_collection(COLLECTION_NAME)
    except Exception:
        pass

    collection = client.create_collection(
        name=COLLECTION_NAME,
        embedding_function=openai_ef,
        metadata={"hnsw:space": "cosine"},
    )

    # Ajout par lots : l'embedding OpenAI est appelé automatiquement par Chroma
    LOT = 100
    for i in range(0, len(documents), LOT):
        collection.add(
            documents=documents[i:i + LOT],
            metadatas=metadatas[i:i + LOT],
            ids=ids[i:i + LOT],
        )
        print(f"  indexé {min(i + LOT, len(documents))}/{len(documents)}")

    print(f"\n[OK] Index construit : {collection.count()} morceaux dans « {COLLECTION_NAME} ».")
    print(f"     Base stockée dans : {CHROMA_DIR}")


if __name__ == "__main__":
    construire_index()
